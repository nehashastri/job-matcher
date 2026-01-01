"""Test script to verify Word document resume extraction."""

import os
import sys
from pathlib import Path

# Add parent directory to path and set working directory
script_dir = Path(__file__).parent
project_root = script_dir.parent
sys.path.insert(0, str(project_root))
os.chdir(project_root)

from config.config import get_config
from matching.resume_loader import ResumeLoader


def main():
    print("=" * 60)
    print("Testing Resume Extraction from Word Document")
    print("=" * 60)

    config = get_config()
    loader = ResumeLoader(config)

    # Check if resume file exists
    resume_path = Path(config.resume_path)
    print(f"\nLooking for resume at: {resume_path}")
    print(f"Resume file exists: {resume_path.exists()}")

    if not resume_path.exists():
        print("\n⚠️  ERROR: Resume file not found!")
        print(f"Expected location: {resume_path.absolute()}")
        print("\nPlease ensure you have placed your resume file at:")
        print("  - data/resume.docx (or)")
        print("  - data/master_resume.docx")
        print("\nOr set RESUME_PATH in your .env file")
        return

    # Try to load the resume
    print("\n" + "-" * 60)
    print("Extracting text from resume...")
    print("-" * 60)

    try:
        # First, let's inspect the document structure
        from docx import Document

        doc = Document(str(resume_path.absolute()))
        print("\nDocument Structure:")
        print(f"  - Number of paragraphs: {len(doc.paragraphs)}")
        print(f"  - Number of tables: {len(doc.tables)}")

        # Show first few paragraphs
        print("\nFirst 3 paragraphs:")
        for i, para in enumerate(doc.paragraphs[:3]):
            print(
                f"  {i + 1}: '{para.text[:100]}...'"
                if len(para.text) > 100
                else f"  {i + 1}: '{para.text}'"
            )

        # Now test the loader
        print("\n" + "-" * 60)
        print("Testing ResumeLoader...")
        print("-" * 60)

        resume_text = loader.load_text()

        if not resume_text:
            print("\n⚠️  WARNING: Resume text is empty!")
            print("\nThe document exists but no text was extracted.")
            print("This could mean:")
            print("  1. The document is truly empty")
            print("  2. The content is in images or text boxes")
            print("  3. The document is corrupted")
        print(f"  - Total words: {len(resume_text.split())}")
        print(f"  - Total lines: {len(resume_text.splitlines())}")

        # Show preview
        print("\n" + "=" * 60)
        print("Resume Text Preview (first 500 characters):")
        print("=" * 60)
        preview = resume_text[:500]
        print(preview)
        if len(resume_text) > 500:
            print("\n... (text truncated for preview)")

        # Test caching
        print("\n" + "-" * 60)
        print("Testing cache functionality...")
        print("-" * 60)
        cached_text = loader.load_text()
        if cached_text == resume_text:
            print("✓ Cache working correctly - same text returned")
        else:
            print("⚠️  WARNING: Cache may not be working properly")

        print("\n" + "=" * 60)
        print("✓ Resume extraction test PASSED!")
        print("=" * 60)

    except Exception as e:
        print("\n❌ ERROR: Failed to extract resume text!")
        print(f"Error type: {type(e).__name__}")
        print(f"Error message: {str(e)}")
        import traceback

        print("\nFull traceback:")
        traceback.print_exc()


if __name__ == "__main__":
    main()

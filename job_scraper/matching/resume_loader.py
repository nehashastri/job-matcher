"""
Resume loader for extracting full resume text.

Reads Word (.docx) or plain text resumes without truncation and caches the result per
path to avoid repeated disk I/O. Defaults to `config.resume_path` when no path
is provided. Uses python-docx for Word document extraction and UTF-8 fallback for text files.

Features:
- Loads resume from .docx or .txt
- Caches loaded resume text for performance
- Uses config.resume_path if no path provided
"""

from __future__ import annotations

from pathlib import Path

from config.config import Config, get_config
from docx import Document


class ResumeLoader:
    """
    Loads resume text from .docx or .txt files, caches results for performance.
    Attributes:
        config (Config): Configuration instance
        logger: Logger instance
        _cache (dict[str, str]): Cache of loaded resume texts
    """

    def __init__(self, config: Config | None = None, logger=None):
        """
        Initialize ResumeLoader.
        Args:
            config (Config | None): Configuration instance
            logger: Logger instance
        """
        self.config = config or get_config()
        self.logger = logger
        self._cache: dict[str, str] = {}

    def load_text(self, path: str | Path | None = None) -> str:
        """
        Load full resume text from the provided path or default config path.
        Returns an empty string if the file is missing or unreadable. Does not
        truncate content; callers are responsible for any prompt-length control.
        Args:
            path (str | Path | None): Path to resume file
        Returns:
            str: Resume text
        """
        target_path = Path(path) if path else Path(self.config.resume_path)

        if target_path.exists() and target_path.is_file():
            cache_key = str(target_path.resolve())
            if cache_key in self._cache:
                return self._cache[cache_key]

            try:
                if target_path.suffix.lower() == ".docx":
                    doc = Document(str(target_path))
                    # Extract text from paragraphs
                    text_parts = [
                        paragraph.text
                        for paragraph in doc.paragraphs
                        if paragraph.text.strip()
                    ]

                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " ".join(
                                cell.text.strip()
                                for cell in row.cells
                                if cell.text.strip()
                            )
                            if row_text:
                                text_parts.append(row_text)

                    text = "\n".join(text_parts)
                else:
                    text = target_path.read_text(encoding="utf-8", errors="ignore")

                if self.logger:
                    self.logger.info(f"Loaded resume text from {target_path}")
                self._cache[cache_key] = text
                return text
            except Exception as exc:  # pragma: no cover - defensive
                if self.logger:
                    self.logger.warning(
                        f"Could not read resume at {target_path}: {exc}"
                    )
                return ""

        if self.logger:
            self.logger.warning(
                f"Resume not found at {target_path}; provide RESUME_PATH."
            )
        return ""
